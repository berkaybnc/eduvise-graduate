import sys
import docx

def update_docx(file_path):
    doc = docx.Document(file_path)
    
    # 1. Modular Monolith Architecture
    for i, p in enumerate(doc.paragraphs):
        if "4.1.2. Modular Monolith Architecture" in p.text:
            # The next paragraphs should be the ones to replace
            # We'll replace the first one and clear the next ones until 4.2
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = "The EduVise platform adopts a modular monolith architectural approach to enhance scalability and maintainability while keeping the deployment complexity low for the current phase. Each core functionality is encapsulated within an independent modular router (e.g., User Service, Course Service, Assessment Service) within a single unified backend application. This approach reduces system coupling and allows individual domains to evolve independently. While the system operates as a single deployable unit currently, the strict modular boundaries ensure that it can easily be transitioned into a microservices architecture in the future as user demand increases."
            if i + 2 < len(doc.paragraphs) and "The primary services include" in doc.paragraphs[i+2].text:
                doc.paragraphs[i+2].text = ""
            if i + 3 < len(doc.paragraphs) and "This architectural decision" in doc.paragraphs[i+3].text:
                doc.paragraphs[i+3].text = ""
            break

    # 2. Database Design
    for i, p in enumerate(doc.paragraphs):
        if "4.3. Database Design" in p.text:
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = "The database layer is designed to support structured schema data requirements of the platform rapidly and reliably. For the current prototype phase of the project, a unified database approach is adopted to optimize development speed and reduce infrastructure complexity.\n\nSQLite is utilized as the primary relational database system for storing all platform data, including user accounts, roles, course content, assessment questions, knowledge graph representations, and transactional records. This unified design supports efficient data access patterns, maintains strict data integrity through foreign keys, and provides a highly portable database solution without the need for external database servers. Future phases of the project may migrate specific high-throughput tables to PostgreSQL and introduce Redis for caching, but the current SQLite implementation adequately handles the system's performance requirements."
            # clear the list paragraphs
            j = i + 2
            while j < len(doc.paragraphs) and not doc.paragraphs[j].text.startswith("4.3.1"):
                doc.paragraphs[j].text = ""
                j += 1
            break

    # 3. API Design
    for i, p in enumerate(doc.paragraphs):
        if "4.4. API Design" in p.text:
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = "The EduVise platform adopts a RESTful API design to enable communication between the client applications and backend services. The backend, powered by FastAPI, acts as a centralized API gateway and handles request routing, authentication, and security enforcement.\n\nPlanned API functionalities include user authentication endpoints, course and content retrieval, assessment submission, and AI-driven roadmap generation requests. Communication between the backend application and the external AI Engine (Large Language Models) is conducted securely via standard HTTPS REST calls, ensuring low-latency and reliable data exchange without the overhead of maintaining a separate internal AI microservice. Real-time updates are handled seamlessly through optimized API polling mechanisms in the current iteration."
            j = i + 2
            while j < len(doc.paragraphs) and not doc.paragraphs[j].text.startswith("4.4.1"):
                doc.paragraphs[j].text = ""
                j += 1
            break

    # 4. Implementation Details
    for i, p in enumerate(doc.paragraphs):
        if "4.5. Implementation Details" in p.text:
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = "In the current phase, the system is implemented using isolated virtual environments (e.g., Python venv) and Node.js package managers to ensure environment consistency across development setups. Configuration management is securely handled through environment variables (.env files), enabling the platform to seamlessly adapt to different local or cloud deployment environments.\n\nError handling and logging mechanisms are built directly into the FastAPI middleware to ensure system reliability and ease of maintenance. Advanced containerization (such as Docker) is planned for the subsequent production deployment phase to further standardize the hosting infrastructure."
            j = i + 2
            while j < len(doc.paragraphs) and not doc.paragraphs[j].text.startswith("4.5.1") and not doc.paragraphs[j].text.startswith("This structured"):
                doc.paragraphs[j].text = ""
                j += 1
            break

    doc.save('graduate_1_final_updated.docx')

update_docx('graduate_1_final_with_images.docx')
print('Success!')
